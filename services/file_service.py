"""
File Service Module - Phase 6.2 Enhanced

This module handles file processing operations for the post generation workflow.
It provides comprehensive text extraction capabilities from multiple file formats
with robust error handling and logging.

Supported formats:
- Plain text (.txt, .md)
- Microsoft Word (.docx)
- PDF documents (.pdf)
- Excel spreadsheets (.xlsx) for post history

Architecture:
- Type-safe interfaces with comprehensive validation
- Structured logging for debugging and monitoring
- Graceful error handling with specific exception types
- Memory-efficient processing for large files
"""

import logging
import io
from typing import List, Optional, Union, Any
from pathlib import Path

# Configure logging for this module
logger = logging.getLogger(__name__)


# Custom Exception Classes
class FileProcessingError(Exception):
    """Base exception for file processing errors."""
    pass


class UnsupportedFileTypeError(FileProcessingError):
    """Exception raised when file type is not supported."""
    pass


class FileReadError(FileProcessingError):
    """Exception raised when file cannot be read."""
    pass


class FileValidationError(FileProcessingError):
    """Exception raised when file validation fails."""
    pass


def extract_text_from_uploads(uploaded_files: List[Any]) -> str:
    """
    Extract and concatenate text content from multiple uploaded files.
    
    This function processes a list of uploaded files, extracts text content from each
    supported file type, and concatenates the results with proper separation.
    
    Args:
        uploaded_files (List[Any]): List of Streamlit UploadedFile objects containing
                                   the files to process. Each file should have 'name' and 'read' attributes.
        
    Returns:
        str: Concatenated text content from all files, separated by double newlines.
             Returns empty string if no files provided or no text extracted.
        
    Raises:
        FileProcessingError: If file processing fails for any reason
        UnsupportedFileTypeError: If any file has an unsupported format
        FileReadError: If any file cannot be read
        FileValidationError: If file validation fails
        
    Example:
        >>> files = [txt_file, docx_file, pdf_file]
        >>> combined_text = extract_text_from_uploads(files)
        >>> print(len(combined_text))
        1542
        
    Note:
        - Files are processed in the order provided
        - Empty files are skipped but logged as warnings
        - Large files are processed with memory optimization
        - Text extraction preserves formatting where possible
    """
    logger.info(f"Starting text extraction from {len(uploaded_files) if uploaded_files else 0} files")
    
    # Input validation
    if not uploaded_files:
        logger.warning("No files provided for text extraction")
        return ""
    
    if not isinstance(uploaded_files, (list, tuple)):
        raise FileValidationError("uploaded_files must be a list or tuple")
    
    extracted_texts = []
    
    for i, file in enumerate(uploaded_files):
        try:
            # Validate file object
            if not hasattr(file, 'name') or not hasattr(file, 'read'):
                raise FileValidationError(f"File {i+1} is not a valid file object")
            
            filename = file.name
            logger.debug(f"Processing file {i+1}/{len(uploaded_files)}: {filename}")
            
            # Extract file extension safely
            if '.' not in filename:
                raise UnsupportedFileTypeError(f"File '{filename}' has no extension")
            
            file_extension = filename.lower().split('.')[-1]
            
            # Route to appropriate processor based on file type
            if file_extension in ['txt', 'md']:
                text = _read_txt_file(file)
            elif file_extension == 'docx':
                text = _read_docx_file(file)
            elif file_extension == 'pdf':
                text = _read_pdf_file(file)
            else:
                raise UnsupportedFileTypeError(f"Unsupported file type: .{file_extension}")
            
            # Validate extracted text
            if text is None:
                logger.warning(f"No text extracted from {filename}")
                continue
            
            text = text.strip()
            if not text:
                logger.warning(f"File {filename} appears to be empty after processing")
                continue
            
            extracted_texts.append(text)
            logger.info(f"Successfully extracted {len(text)} characters from {filename}")
            
        except (UnsupportedFileTypeError, FileValidationError):
            # Re-raise validation and type errors as-is
            raise
        except Exception as e:
            logger.error(f"Failed to process file {filename}: {str(e)}")
            raise FileReadError(f"Failed to read file '{filename}': {str(e)}") from e
    
    if not extracted_texts:
        logger.warning("No text content was extracted from any files")
        return ""
    
    # Combine all extracted texts
    combined_text = "\n\n".join(extracted_texts)
    logger.info(f"Text extraction completed: {len(extracted_texts)} files, {len(combined_text)} total characters")
    
    return combined_text


def extract_posts_from_history(history_file: Any) -> List[str]:
    """
    Extract post text from Excel history file with comprehensive validation.
    
    This function reads an Excel file containing post history data and extracts
    the post text content for use in learning posting patterns and style.
    
    Args:
        history_file (Any): Streamlit UploadedFile object containing Excel data (.xlsx).
                           Must have 'read' method and 'name' attribute.
        
    Returns:
        List[str]: List of post text strings extracted from the 'Post Text' column.
                  Empty list if file is empty or no valid posts found.
                  
    Raises:
        FileValidationError: If file object is invalid or missing required attributes
        FileReadError: If Excel file cannot be read or is corrupted
        KeyError: If required 'Post Text' column is not found
        FileProcessingError: If processing fails for any other reason
        
    Expected Excel Format:
        - Must contain a column named 'Post Text'
        - Optional columns: 'Platform', 'Date', 'Engagement', etc.
        - First row should contain headers
        
    Example:
        >>> history_posts = extract_posts_from_history(excel_file)
        >>> print(f"Found {len(history_posts)} historical posts")
        Found 15 historical posts
        
    Note:
        - Empty cells in 'Post Text' column are filtered out
        - Very short posts (< 10 characters) are logged as warnings
        - Large Excel files are processed efficiently with pandas
        - File is automatically closed after processing
    """
    import pandas as pd
    
    logger.info(f"Starting post history extraction from {getattr(history_file, 'name', 'unknown file')}")
    
    # Input validation
    if not history_file:
        raise FileValidationError("History file cannot be None")
    
    if not hasattr(history_file, 'read'):
        raise FileValidationError("History file must have a 'read' method")
    
    try:
        # Read the Excel file content
        logger.debug("Reading Excel file content")
        file_content = history_file.read()
        
        if not file_content:
            logger.warning("History file appears to be empty")
            return []
        
        # Parse Excel content using pandas
        logger.debug("Parsing Excel content with pandas")
        df = pd.read_excel(io.BytesIO(file_content))
        
        # Handle empty dataframe
        if df.empty:
            logger.warning("Excel file contains no data rows")
            return []
        
        logger.info(f"Excel file loaded: {len(df)} rows, {len(df.columns)} columns")
        logger.debug(f"Available columns: {list(df.columns)}")
        
        # Check if required column exists
        if 'Post Text' not in df.columns:
            available_cols = ', '.join(df.columns)
            error_msg = f"Required column 'Post Text' not found. Available columns: {available_cols}"
            logger.error(error_msg)
            raise KeyError(error_msg)
        
        # Extract and validate post text
        logger.debug("Extracting post text content")
        raw_posts = df['Post Text'].tolist()
        
        # Filter and validate posts
        valid_posts = []
        for i, post in enumerate(raw_posts):
            # Skip None/NaN values
            if pd.isna(post) or post is None:
                logger.debug(f"Skipping empty post at row {i+1}")
                continue
            
            # Convert to string and strip whitespace
            post_text = str(post).strip()
            
            # Skip empty posts
            if not post_text:
                logger.debug(f"Skipping whitespace-only post at row {i+1}")
                continue
            
            # Warn about very short posts
            if len(post_text) < 10:
                logger.warning(f"Very short post at row {i+1}: '{post_text[:20]}...'")
            
            valid_posts.append(post_text)
        
        logger.info(f"Post history extraction completed: {len(valid_posts)} valid posts from {len(raw_posts)} total rows")
        return valid_posts
        
    except pd.errors.EmptyDataError:
        logger.warning("Excel file contains no data")
        return []
    except pd.errors.ParserError as e:
        logger.error(f"Failed to parse Excel file: {str(e)}")
        raise FileReadError(f"Excel file appears to be corrupted or invalid: {str(e)}") from e
    except Exception as e:
        logger.error(f"Unexpected error during post history extraction: {str(e)}")
        raise FileProcessingError(f"Failed to extract post history: {str(e)}") from e


def _read_txt_file(file: Any) -> str:
    """
    Read and decode plain text files with comprehensive encoding support.
    
    Handles text files (.txt, .md) with automatic encoding detection and fallback.
    Supports UTF-8, Latin-1, and ASCII encodings with graceful degradation.
    
    Args:
        file (Any): File object with read() method
        
    Returns:
        str: Decoded text content from the file
        
    Raises:
        FileReadError: If file cannot be read or decoded
        
    Note:
        - Attempts UTF-8 decoding first (recommended)
        - Falls back to Latin-1 if UTF-8 fails
        - Uses ASCII with error ignoring as last resort
    """
    filename = getattr(file, 'name', 'unknown')
    logger.debug(f"Reading text file: {filename}")
    
    try:
        content = file.read()
        
        if not content:
            logger.warning(f"Text file {filename} is empty")
            return ""
        
        # Handle both bytes and string content
        if isinstance(content, bytes):
            logger.debug(f"Decoding bytes content from {filename}")
            
            # Try UTF-8 first (recommended encoding)
            try:
                decoded = content.decode('utf-8')
                logger.debug(f"Successfully decoded {filename} as UTF-8")
                return decoded
            except UnicodeDecodeError as e:
                logger.warning(f"UTF-8 decoding failed for {filename}: {str(e)}")
                
                # Fall back to Latin-1
                try:
                    decoded = content.decode('latin-1')
                    logger.warning(f"Using Latin-1 fallback encoding for {filename}")
                    return decoded
                except UnicodeDecodeError as e:
                    logger.warning(f"Latin-1 decoding failed for {filename}: {str(e)}")
                    
                    # Last resort: ASCII with error handling
                    try:
                        decoded = content.decode('ascii', errors='ignore')
                        logger.warning(f"Using ASCII with error ignoring for {filename} (some characters may be lost)")
                        return decoded
                    except Exception as e:
                        logger.error(f"All encoding attempts failed for {filename}: {str(e)}")
                        raise FileReadError(f"Cannot decode text file '{filename}': {str(e)}") from e
        
        elif isinstance(content, str):
            logger.debug(f"Content from {filename} is already a string")
            return content
        
        else:
            logger.error(f"Unexpected content type from {filename}: {type(content)}")
            raise FileReadError(f"Unexpected content type from '{filename}': {type(content)}")
            
    except Exception as e:
        if isinstance(e, FileReadError):
            raise
        logger.error(f"Failed to read text file {filename}: {str(e)}")
        raise FileReadError(f"Failed to read text file '{filename}': {str(e)}") from e


def _read_docx_file(file: Any) -> str:
    """
    Extract text content from Microsoft Word documents (.docx).
    
    Comprehensively extracts text from paragraphs, tables, headers, and footers
    using the python-docx library with robust error handling.
    
    Args:
        file (Any): File object containing DOCX data
        
    Returns:
        str: Extracted text content with preserved structure
        
    Raises:
        FileReadError: If DOCX file cannot be read or processed
        ImportError: If python-docx library is not available
        
    Note:
        - Extracts text from paragraphs, tables, headers, and footers
        - Preserves basic document structure with line breaks
        - Table content is formatted with pipe separators
        - Handles both simple and complex document layouts
    """
    filename = getattr(file, 'name', 'unknown')
    logger.debug(f"Reading DOCX file: {filename}")
    
    try:
        from docx import Document
    except ImportError as e:
        logger.error("python-docx library not available")
        raise ImportError("python-docx library is required for DOCX file processing. Install with: pip install python-docx") from e
    
    try:
        # Read the file content into BytesIO
        logger.debug(f"Reading file content from {filename}")
        file_content = file.read()
        
        if not file_content:
            logger.warning(f"DOCX file {filename} is empty")
            return ""
        
        # Create document stream
        doc_stream = io.BytesIO(file_content)
        
        # Load the document
        logger.debug(f"Loading DOCX document: {filename}")
        doc = Document(doc_stream)
        
        content_parts = []
        
        # Extract text from paragraphs
        logger.debug(f"Extracting paragraphs from {filename}")
        paragraph_count = 0
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                content_parts.append(text)
                paragraph_count += 1
        
        logger.debug(f"Extracted {paragraph_count} paragraphs from {filename}")
        
        # Extract text from tables
        logger.debug(f"Extracting tables from {filename}")
        table_count = 0
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                
                if row_cells:
                    table_content.append(" | ".join(row_cells))
            
            if table_content:
                content_parts.extend(table_content)
                table_count += 1
        
        logger.debug(f"Extracted {table_count} tables from {filename}")
        
        # Extract text from headers and footers
        logger.debug(f"Extracting headers and footers from {filename}")
        header_footer_count = 0
        for section in doc.sections:
            # Headers
            if section.header:
                for paragraph in section.header.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        content_parts.append(f"[Header: {text}]")
                        header_footer_count += 1
            
            # Footers
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        content_parts.append(f"[Footer: {text}]")
                        header_footer_count += 1
        
        if header_footer_count > 0:
            logger.debug(f"Extracted {header_footer_count} header/footer elements from {filename}")
        
        # Combine all content
        final_content = "\n".join(content_parts)
        
        logger.info(f"DOCX extraction completed for {filename}: {len(content_parts)} content blocks, {len(final_content)} characters")
        return final_content
        
    except Exception as e:
        logger.error(f"Failed to process DOCX file {filename}: {str(e)}")
        raise FileReadError(f"Failed to read DOCX file '{filename}': {str(e)}") from e


def _read_pdf_file(file: Any) -> str:
    """
    Extract text content from PDF documents using PyMuPDF.
    
    Comprehensively extracts text from all pages of PDF documents with
    support for various PDF formats and layouts.
    
    Args:
        file (Any): File object containing PDF data
        
    Returns:
        str: Extracted text content from all pages
        
    Raises:
        FileReadError: If PDF file cannot be read or processed
        ImportError: If PyMuPDF library is not available
        Exception: If PDF is password protected or corrupted
        
    Note:
        - Processes all pages in the document
        - Handles text extraction from complex layouts
        - Automatically detects and handles password-protected PDFs
        - Preserves page structure with appropriate separators
        - Memory-efficient for large PDF files
    """
    filename = getattr(file, 'name', 'unknown')
    logger.debug(f"Reading PDF file: {filename}")
    
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        logger.error("PyMuPDF library not available")
        raise ImportError("PyMuPDF library is required for PDF file processing. Install with: pip install PyMuPDF") from e
    
    pdf_document = None
    try:
        # Read the file content
        logger.debug(f"Reading file content from {filename}")
        file_content = file.read()
        
        if not file_content:
            logger.warning(f"PDF file {filename} is empty")
            return ""
        
        # Open the PDF from bytes
        logger.debug(f"Opening PDF document: {filename}")
        pdf_document = fitz.open(stream=file_content, filetype="pdf")
        
        # Check if the document requires a password
        if pdf_document.needs_pass:
            logger.error(f"PDF file {filename} is password protected")
            raise FileReadError(f"PDF file '{filename}' is password protected and cannot be processed")
        
        # Get document metadata
        page_count = pdf_document.page_count
        logger.info(f"Processing PDF {filename}: {page_count} pages")
        
        if page_count == 0:
            logger.warning(f"PDF file {filename} contains no pages")
            return ""
        
        # Extract text from all pages
        text_content = []
        pages_with_text = 0
        
        for page_num in range(page_count):
            try:
                logger.debug(f"Extracting text from page {page_num + 1}/{page_count} of {filename}")
                page = pdf_document[page_num]
                
                # Extract text using different methods for better coverage
                text = page.get_text()
                
                # Alternative text extraction if primary method yields little content
                if len(text.strip()) < 50:
                    logger.debug(f"Primary text extraction yielded little content for page {page_num + 1}, trying alternative method")
                    try:
                        # Try getting text blocks
                        blocks = page.get_text("blocks")
                        alt_text = "\n".join([block[4] for block in blocks if len(block) > 4 and block[4].strip()])
                        if len(alt_text.strip()) > len(text.strip()):
                            text = alt_text
                    except Exception as e:
                        logger.debug(f"Alternative text extraction failed for page {page_num + 1}: {str(e)}")
                
                if text.strip():
                    text_content.append(text.strip())
                    pages_with_text += 1
                else:
                    logger.debug(f"No text found on page {page_num + 1} of {filename}")
                    
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1} of {filename}: {str(e)}")
                continue
        
        logger.info(f"PDF extraction completed for {filename}: {pages_with_text}/{page_count} pages with text, {len(text_content)} text blocks")
        
        if not text_content:
            logger.warning(f"No text content extracted from PDF {filename}")
            return ""
        
        # Combine all page content
        final_content = "\n\n".join(text_content)
        logger.info(f"PDF {filename} processing completed: {len(final_content)} total characters")
        
        return final_content
        
    except Exception as e:
        if isinstance(e, FileReadError):
            raise
        logger.error(f"Failed to process PDF file {filename}: {str(e)}")
        raise FileReadError(f"Failed to read PDF file '{filename}': {str(e)}") from e
    
    finally:
        # Ensure PDF document is always closed
        if pdf_document:
            try:
                pdf_document.close()
                logger.debug(f"Closed PDF document: {filename}")
            except Exception as e:
                logger.warning(f"Failed to close PDF document {filename}: {str(e)}")


# Utility Functions

def get_supported_file_types() -> List[str]:
    """
    Get list of supported file extensions.
    
    Returns:
        List[str]: List of supported file extensions including the dot
    """
    return ['.txt', '.md', '.docx', '.pdf', '.xlsx']


def validate_file_type(filename: str) -> bool:
    """
    Validate if file type is supported.
    
    Args:
        filename (str): Name of the file to validate
        
    Returns:
        bool: True if file type is supported, False otherwise
    """
    if not filename or '.' not in filename:
        return False
    
    extension = '.' + filename.lower().split('.')[-1]
    return extension in get_supported_file_types()


def get_file_type_info() -> dict:
    """
    Get information about supported file types.
    
    Returns:
        dict: Dictionary mapping file extensions to descriptions
    """
    return {
        '.txt': 'Plain text files',
        '.md': 'Markdown files',
        '.docx': 'Microsoft Word documents',
        '.pdf': 'PDF documents',
        '.xlsx': 'Excel spreadsheets (for post history)'
    }